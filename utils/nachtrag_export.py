from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import date


def erstelle_nachtrag_docx(
    baustelle_name: str,
    auftraggeber: str,
    datum: str,
    beschreibung: str,
    ki_text: str
) -> bytes:
    doc = Document()

    # Seitenränder
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # Titel
    titel = doc.add_heading("MEHRKOSTENANZEIGE", 0)
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Metadaten-Tabelle
    tabelle = doc.add_table(rows=4, cols=2)
    tabelle.style = "Table Grid"
    felder = [
        ("Baustelle:", baustelle_name),
        ("Auftraggeber:", auftraggeber),
        ("Datum:", datum or str(date.today())),
        ("Betreff:", beschreibung[:100]),
    ]
    for i, (label, wert) in enumerate(felder):
        tabelle.rows[i].cells[0].text = label
        tabelle.rows[i].cells[1].text = wert
        tabelle.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # KI-generierter Text
    doc.add_heading("KI-Analyse (Gemini)", level=2)

    for zeile in ki_text.split("\n"):
        zeile = zeile.strip()
        if not zeile:
            doc.add_paragraph()
        elif zeile.startswith("**") and zeile.endswith("**"):
            p = doc.add_paragraph(zeile.replace("**", ""))
            p.runs[0].bold = True
        else:
            doc.add_paragraph(zeile)

    doc.add_paragraph()

    # Unterschriftenzeile
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"Datum, Unterschrift Bauleiter")

    # In Bytes umwandeln
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
