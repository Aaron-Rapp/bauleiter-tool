import streamlit as st
from utils.db import get_db_client
import datetime
import html as html_mod
from utils.config import get_config, get_name

def _ki_call(prompt: str) -> str:
    api_key = get_config("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("Kein GEMINI_API_KEY gesetzt.")
    from google import genai
    client = genai.Client(api_key=api_key)
    resp = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return resp.text

try:
    from streamlit_calendar import calendar as st_calendar
    _HAS_CALENDAR = True
except ImportError:
    _HAS_CALENDAR = False

try:
    from streamlit_pdf_viewer import pdf_viewer as _pdf_viewer
    _HAS_PDF = True
except ImportError:
    _HAS_PDF = False

try:
    from streamlit_folium import st_folium as _st_folium
    import folium as _folium
    _HAS_FOLIUM = True
except ImportError:
    _HAS_FOLIUM = False

st.set_page_config(page_title="Projekt", page_icon="🏗️", layout="wide", initial_sidebar_state="collapsed")

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:ital,wght@0,300;0,400;0,500;0,600;0,700;0,800;0,900&display=swap');

html, body, [class*="css"] {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif !important;
}
[data-testid="collapsedControl"] { display: none; }
.stApp { background: #F2F1EF !important; }
.main .block-container { padding: 1.8rem 2.5rem 3rem 2.5rem; max-width: 1200px; }

h1 { font-weight: 800 !important; letter-spacing: -0.04em !important; color: #111111 !important; }
h2, h3 { font-weight: 700 !important; letter-spacing: -0.025em !important; color: #1A1A1A !important; }

/* Tabs */
div[data-testid="stTabs"] div[role="tablist"],
.stTabs [data-baseweb="tab-list"] {
    background: #E9E8E5 !important; border-radius: 10px !important;
    padding: 3px !important; gap: 1px !important; border: none !important;
}
div[data-testid="stTabs"] button[role="tab"],
.stTabs [data-baseweb="tab"] {
    border-radius: 8px !important; font-weight: 500 !important;
    font-size: 0.84rem !important; padding: 7px 13px !important;
    color: #777672 !important; border: none !important;
    background: transparent !important;
}
div[data-testid="stTabs"] button[role="tab"][aria-selected="true"],
.stTabs [aria-selected="true"] {
    background: white !important; color: #1A1A1A !important;
    font-weight: 600 !important; box-shadow: 0 1px 3px rgba(0,0,0,0.08) !important;
}
div[data-testid="stTabs"] div[role="tablist"]::after,
.stTabs [data-baseweb="tab-highlight"] { display: none !important; }
div[data-testid="stTabs"] div[role="tabpanel"],
.stTabs [data-baseweb="tab-panel"] { padding-top: 1.2rem !important; }

/* Primary Buttons */
.stButton > button[kind="primary"] {
    background: #1A1A1A !important; color: white !important;
    border: none !important; border-radius: 8px !important;
    font-weight: 600 !important; font-size: 0.87rem !important;
    padding: 0.55rem 1.3rem !important; letter-spacing: 0.015em !important;
    transition: background 0.15s !important; box-shadow: none !important;
}
.stButton > button[kind="primary"]:hover {
    background: #333 !important; box-shadow: 0 4px 14px rgba(0,0,0,0.16) !important;
}
/* Text in Primary-Buttons IMMER weiß (sonst schwarz auf schwarz) */
.stButton > button[kind="primary"] *,
[data-testid="stBaseButton-primary"] * {
    color: white !important;
}

/* Secondary Buttons */
.stButton > button[kind="secondary"] {
    border-radius: 8px !important; font-weight: 500 !important;
    font-size: 0.87rem !important; border: 1px solid #DDDBD8 !important;
    background: white !important; color: #333 !important;
    transition: border-color 0.15s, color 0.15s !important; box-shadow: none !important;
}
.stButton > button[kind="secondary"]:hover {
    border-color: #F07030 !important; color: #F07030 !important; background: white !important;
}

/* Cards */
[data-testid="stVerticalBlockBorderWrapper"] {
    border-radius: 14px !important; border: 1px solid #E8E6E2 !important;
    box-shadow: 0 1px 2px rgba(0,0,0,0.03), 0 4px 12px rgba(0,0,0,0.04) !important;
    background: white !important; overflow: hidden !important;
}

/* Inputs — orange focus */
.stTextInput > div > div > input,
.stTextArea > div > div > textarea,
.stDateInput > div > div > input {
    border-radius: 8px !important; border: 1px solid #DDDBD8 !important;
    font-family: 'Inter', sans-serif !important; font-size: 0.89rem !important;
    background: #FAFAF9 !important; color: #1A1A1A !important;
    transition: border-color 0.15s !important;
}
.stTextInput > div > div > input:focus,
.stTextArea > div > div > textarea:focus {
    border-color: #F07030 !important;
    box-shadow: 0 0 0 2.5px rgba(240,112,48,0.14) !important;
    background: white !important;
}
.stSelectbox > div > div { border-radius: 8px !important; border: 1px solid #DDDBD8 !important; background: #FAFAF9 !important; }
.stMultiSelect > div > div { border-radius: 8px !important; border: 1px solid #DDDBD8 !important; }
.stRadio label { font-size: 0.88rem !important; color: #444 !important; }

/* Expander */
.streamlit-expanderHeader {
    border-radius: 10px !important; font-weight: 600 !important;
    font-size: 0.9rem !important; background: white !important;
    border: 1px solid #E8E6E2 !important; color: #1A1A1A !important;
}
.streamlit-expanderContent {
    border: 1px solid #E8E6E2 !important; border-top: none !important;
    border-radius: 0 0 10px 10px !important; background: white !important;
}

/* Alerts */
[data-testid="stAlert"] { border-radius: 10px !important; }
.stInfo { background: #F9F8F6 !important; border-left: 3px solid #DDDBD8 !important; }
.stWarning { border-left: 3px solid #F07030 !important; }

/* HR */
hr { border-color: #E8E6E2 !important; margin: 1rem 0 !important; }

/* Caption */
.stCaption, [data-testid="stCaptionContainer"] p { color: #888683 !important; font-size: 0.81rem !important; }

/* Progress */
.stProgress > div > div { border-radius: 999px !important; }

/* Form */
[data-testid="stForm"] {
    background: #F7F6F4 !important; border-radius: 12px !important;
    padding: 1.2rem !important; border: 1px solid #E8E6E2 !important;
}

/* Image */
[data-testid="stImage"] img { border-radius: 10px !important; }

/* Chat */
[data-testid="stChatMessage"] { border-radius: 12px !important; border: 1px solid #EEECe9 !important; }
[data-testid="stChatInput"] > div { border-radius: 12px !important; border: 1px solid #DDDBD8 !important; }

/* Checkbox */
.stCheckbox label span { font-size: 0.88rem !important; }
</style>
""", unsafe_allow_html=True)

# ── DB ────────────────────────────────────────────────────────
if "db" in st.session_state:
    db = st.session_state.db
else:
    db, _m = get_db_client()
    st.session_state.db = db
    st.session_state.db_modus = _m

modus = st.session_state.get("db_modus", "lokal")

# ── Projekt prüfen ────────────────────────────────────────────
if "aktives_projekt" not in st.session_state:
    st.warning("Kein Projekt ausgewählt.")
    if st.button("Zurück zur Übersicht"):
        st.switch_page("app.py")
    st.stop()

pid   = st.session_state.aktives_projekt
pname = st.session_state.get("aktives_projekt_name", "Projekt")
rolle = st.session_state.get("rolle", "Außenstehend")

# ── Projektdaten laden (für Wetter-Widget) ────────────────────
try:
    projekt_daten = db.table("projekte").select("*").eq("id", pid).execute().data
    projekt = projekt_daten[0] if projekt_daten else {}
except Exception:
    projekt = {}

# ── Header ────────────────────────────────────────────────────
col_nav, col_mid, col_logo = st.columns([2, 5, 3])
with col_nav:
    st.markdown("<div style='padding-top:0.5rem;'>", unsafe_allow_html=True)
    if st.button("← Zurück"):
        st.switch_page("app.py")
    st.markdown("</div>", unsafe_allow_html=True)
with col_mid:
    st.markdown(
        f"<h1 style='margin:0.15rem 0; font-size:1.55rem; white-space:nowrap; overflow:hidden;"
        f"text-overflow:ellipsis;'>{html_mod.escape(pname)}</h1>",
        unsafe_allow_html=True
    )
with col_logo:
    st.markdown("""
    <div style='display:flex; align-items:center; gap:10px; justify-content:flex-end; padding-top:0.3rem;'>
      <div style='
        width:36px; height:36px; background:#F07030; border-radius:9px;
        display:flex; align-items:center; justify-content:center; flex-shrink:0;
        box-shadow:0 2px 8px rgba(240,112,48,0.2);
      '>
        <span style='color:white; font-family:Inter,sans-serif; font-weight:900;
                     font-size:18px; letter-spacing:-0.06em;'>R</span>
      </div>
      <div style='font-family:Inter,sans-serif; font-size:22px; font-weight:800;
                  letter-spacing:-0.045em; color:#1A1A1A; line-height:1;'>
        <span style='color:#F07030;'>·</span>APP
      </div>
    </div>
    """, unsafe_allow_html=True)

st.markdown("<hr style='margin:0.3rem 0 0.5rem 0; border-color:#E8E6E2;'>", unsafe_allow_html=True)

# ── Wetter-Widget ─────────────────────────────────────────────
anschrift = projekt.get("anschrift", "").strip()
if anschrift:
    wetter_key = f"wetter_{pid}"
    coords_key = f"coords_{pid}"
    if wetter_key not in st.session_state:
        try:
            from utils.wetter import geocode, get_wetter
            coords = geocode(anschrift)
            st.session_state[coords_key] = coords or False
            if coords:
                w = get_wetter(coords[0], coords[1])
                st.session_state[wetter_key] = w
            else:
                st.session_state[wetter_key] = {"ok": False}
        except Exception:
            st.session_state[wetter_key] = {"ok": False}
            st.session_state[coords_key] = False

    w = st.session_state.get(wetter_key, {})
    if w.get("ok"):
        from utils.wetter import wetter_icon_svg
        _icon = wetter_icon_svg(w.get("code", 0), size=44)
        st.markdown(
            f"""<div style='background:linear-gradient(135deg,#ffffff,#f6f8fa);
                            border:1px solid #E8E6E2; border-radius:13px;
                            padding:13px 20px; margin-bottom:1.1rem;
                            display:flex; align-items:center; gap:1.15rem; flex-wrap:wrap;
                            box-shadow:0 1px 3px rgba(0,0,0,0.04), 0 4px 14px rgba(0,0,0,0.04);'>
                <div style='flex-shrink:0; display:flex; align-items:center;'>{_icon}</div>
                <div style='display:flex; flex-direction:column; line-height:1.12;'>
                    <span style='font-weight:800; font-size:1.5rem; color:#1A1A1A; letter-spacing:-0.03em;'>{w['temp']} °C</span>
                    <span style='font-size:0.84rem; color:#777672;'>{w['beschreibung']}</span>
                </div>
                <div style='height:36px; width:1px; background:#E8E6E2; margin:0 0.2rem;'></div>
                <div style='display:flex; gap:1.5rem;'>
                    <div style='display:flex; flex-direction:column; line-height:1.2;'>
                        <span style='color:#A8A6A2; font-size:0.68rem; text-transform:uppercase; letter-spacing:0.06em;'>Wind</span>
                        <span style='font-weight:600; font-size:0.92rem; color:#444441;'>{w['wind']} km/h</span>
                    </div>
                    <div style='display:flex; flex-direction:column; line-height:1.2;'>
                        <span style='color:#A8A6A2; font-size:0.68rem; text-transform:uppercase; letter-spacing:0.06em;'>Niederschlag</span>
                        <span style='font-weight:600; font-size:0.92rem; color:#444441;'>{w['niederschlag']} mm</span>
                    </div>
                </div>
                <span style='font-size:0.76rem; color:#AAAAAA; margin-left:auto;
                             white-space:nowrap; overflow:hidden; text-overflow:ellipsis;
                             max-width:240px;'>{html_mod.escape(anschrift)}</span>
            </div>""",
            unsafe_allow_html=True
        )

# ── Karte ─────────────────────────────────────────────────────
if _HAS_FOLIUM and anschrift:
    coords_key = f"coords_{pid}"
    if coords_key not in st.session_state:
        try:
            from utils.wetter import geocode
            st.session_state[coords_key] = geocode(anschrift) or False
        except Exception:
            st.session_state[coords_key] = False
    coords = st.session_state.get(coords_key, False)
    if coords:
        with st.expander("Baustellenstandort anzeigen"):
            m = _folium.Map(location=coords, zoom_start=16, tiles="CartoDB positron")
            _folium.Marker(
                location=coords,
                tooltip=anschrift,
                popup=_folium.Popup(f"<b>{html_mod.escape(pname)}</b><br>{html_mod.escape(anschrift)}", max_width=250),
                icon=_folium.Icon(color="orange", icon="home", prefix="fa")
            ).add_to(m)
            _st_folium(m, use_container_width=True, height=320, returned_objects=[])

@st.dialog("Foto", width="large")
def _zeige_foto_gross(url: str, name: str):
    st.image(url, use_container_width=True)
    st.caption(name)

# ── 8 Tabs ────────────────────────────────────────────────────
tab_todos, tab_kalender, tab_chat, tab_vertraege, tab_plaene, tab_bilder, tab_vob, tab_bericht = st.tabs([
    "✅ Aufgaben", "📅 Kalender", "🤖 KI-Assistent", "📄 Verträge", "📐 Pläne", "📷 Bilder", "⚖️ VOB", "📋 Tagesbericht"
])

# ─────────────────────────────────────────────────────────────
# HELPER
# ─────────────────────────────────────────────────────────────
def upload_datei(datei_bytes: bytes, pfad: str, content_type: str = "application/octet-stream") -> str:
    try:
        db.storage.from_("bauleiter-dateien").upload(
            pfad, datei_bytes, {"content-type": content_type, "upsert": "true"}
        )
        return db.storage.from_("bauleiter-dateien").get_public_url(pfad)
    except Exception:
        import base64 as _b64fb
        return f"data:{content_type};base64," + _b64fb.b64encode(datei_bytes).decode()

def get_unterordner(kategorie: str) -> list:
    try:
        rows = db.table("dateien").select("unterordner").eq("projekt_id", pid).eq("kategorie", kategorie).execute().data
        ordner = list({r["unterordner"] for r in rows if r["unterordner"]})
        return sorted(ordner) if ordner else ["Allgemein"]
    except Exception:
        return ["Allgemein"]

def speichere_unterordner(kategorie: str, ordner_name: str):
    try:
        vorhanden = db.table("dateien").select("id").eq("projekt_id", pid).eq("kategorie", kategorie).eq("unterordner", ordner_name).execute().data
        if not vorhanden:
            db.table("dateien").insert({
                "projekt_id": pid, "kategorie": kategorie,
                "unterordner": ordner_name, "datei_name": ".ordner", "datei_url": ""
            }).execute()
    except Exception:
        pass

# ─────────────────────────────────────────────────────────────
# SCHRIFTVERKEHR-HELPER
# ─────────────────────────────────────────────────────────────
import json as _json

def get_naechste_nummer(typ: str) -> int:
    try:
        rows = db.table("schriftverkehr").select("nummer").eq("projekt_id", pid).eq("typ", typ).execute().data
        return max((r["nummer"] or 0) for r in rows) + 1 if rows else 1
    except Exception:
        return 1

def speichere_dokument(typ: str, nummer: int, titel: str, inhalt: str, meta: dict) -> bool:
    try:
        db.table("schriftverkehr").insert({
            "projekt_id": pid, "typ": typ, "nummer": nummer,
            "titel": titel, "inhalt": inhalt,
            "meta": _json.dumps(meta, ensure_ascii=False),
        }).execute()
        return True
    except Exception as e:
        st.error(f"Fehler beim Speichern: {e}")
        return False

def lade_dokumente(typ: str) -> list:
    try:
        rows = db.table("schriftverkehr").select("*").eq("projekt_id", pid).eq("typ", typ).order("nummer").execute().data
        return rows or []
    except Exception:
        return []

# ═════════════════════════════════════════════════════════════
# TAB 1: BILDER
# ═════════════════════════════════════════════════════════════
with tab_bilder:
    st.markdown("### Baustellen-Fotos")

    ordner_liste = get_unterordner("Bilder")
    if "Allgemein" not in ordner_liste:
        ordner_liste = ["Allgemein"] + ordner_liste

    col_ord, col_neu = st.columns([3, 2])
    with col_ord:
        akt_ordner = st.selectbox("Ordner", ordner_liste, key="bild_ord")
    with col_neu:
        with st.popover("Neuer Unterordner"):
            neu = st.text_input("Ordner-Name", key="neu_bild_inp")
            if st.button("Anlegen", key="neu_bild_btn") and neu.strip():
                speichere_unterordner("Bilder", neu.strip())
                st.success(f"Ordner '{neu}' angelegt!")
                st.rerun()

    col_cam, col_up = st.columns(2)
    with col_cam:
        foto = st.camera_input("Direkt fotografieren")
    with col_up:
        upload = st.file_uploader("Oder Bild hochladen", type=["jpg","jpeg","png","webp"], key="bild_up")

    bild_datei = foto or upload
    if bild_datei:
        if st.button("Bild speichern", type="primary", key="bild_save"):
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            _bild_bytes = bild_datei.getvalue()
            if modus == "supabase":
                pfad = f"{pid}/Bilder/{akt_ordner}/{ts}.jpg"
                url = upload_datei(_bild_bytes, pfad, "image/jpeg")
            else:
                import base64
                url = "data:image/jpeg;base64," + base64.b64encode(_bild_bytes).decode()
            if url:
                db.table("dateien").insert({
                    "projekt_id": pid, "kategorie": "Bilder",
                    "unterordner": akt_ordner, "datei_name": f"{ts}.jpg", "datei_url": url
                }).execute()
                st.success("Bild gespeichert!")
                st.rerun()

    st.markdown("---")
    try:
        bilder = db.table("dateien").select("*").eq("projekt_id", pid).eq("kategorie", "Bilder").eq("unterordner", akt_ordner).neq("datei_name", ".ordner").order("erstellt_am", desc=True).execute().data
        if bilder:
            st.caption(f"{len(bilder)} Foto{'s' if len(bilder) != 1 else ''} in '{akt_ordner}'")
            cols = st.columns(2)
            for i, b in enumerate(bilder):
                with cols[i % 2]:
                    with st.container(border=True):
                        try:
                            st.image(b["datei_url"], use_container_width=True)
                        except Exception:
                            st.markdown(f"📷 {b['datei_name']}")
                        datum = b["erstellt_am"][:10] if b.get("erstellt_am") else ""
                        name = b.get("datei_name", "")
                        c1, c2, c3 = st.columns([4, 2, 2])
                        with c1:
                            st.caption(f"{datum}")
                        with c2:
                            if st.button("Vollbild", key=f"voll_{b['id']}", use_container_width=True):
                                _zeige_foto_gross(b["datei_url"], f"{name} · {datum}")
                        with c3:
                            if st.button("Löschen", key=f"del_bild_{b['id']}", use_container_width=True):
                                db.table("dateien").delete().eq("id", b["id"]).execute()
                                st.rerun()
        else:
            st.markdown(f"""
            <div style='text-align:center; padding:2rem; color:#AAAAAA;'>
                <svg width="40" height="40" viewBox="0 0 24 24" fill="none" stroke="#cbd5e1" stroke-width="1.5" style="margin-bottom:0.5rem; display:block; margin-left:auto; margin-right:auto;">
                    <rect x="3" y="3" width="18" height="18" rx="2"/>
                    <circle cx="8.5" cy="8.5" r="1.5"/>
                    <polyline points="21 15 16 10 5 21"/>
                </svg>
                <p style='font-size:0.9rem; margin:0;'>Noch keine Fotos in <b>'{akt_ordner}'</b></p>
            </div>""", unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Fehler: {e}")

# ═════════════════════════════════════════════════════════════
# TAB 2: PLÄNE
# ═════════════════════════════════════════════════════════════
with tab_plaene:
    st.markdown("### Baupläne")

    plan_ordner_liste = get_unterordner("Plaene")
    if "Allgemein" not in plan_ordner_liste:
        plan_ordner_liste = ["Allgemein"] + plan_ordner_liste

    col_ord2, col_neu2 = st.columns([3, 2])
    with col_ord2:
        plan_ord = st.selectbox("Ordner", plan_ordner_liste, key="plan_ord")
    with col_neu2:
        with st.popover("Neuer Unterordner"):
            neu2 = st.text_input("Ordner-Name", key="neu_plan_inp")
            if st.button("Anlegen", key="neu_plan_btn") and neu2.strip():
                speichere_unterordner("Plaene", neu2.strip())
                st.success(f"Ordner '{neu2}' angelegt!")
                st.rerun()

    with st.form("plan_form", clear_on_submit=True):
        plan_up = st.file_uploader("Plan hochladen (PDF, DWG, PNG, JPG)", type=["pdf","dwg","png","jpg","jpeg"])
        if st.form_submit_button("Plan speichern", type="primary") and plan_up:
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            pfad = f"{pid}/Plaene/{plan_ord}/{ts}_{plan_up.name}"
            if modus == "supabase":
                url = upload_datei(plan_up.getvalue(), pfad)
            else:
                import base64 as _b64mod
                _mime = "application/pdf" if plan_up.name.lower().endswith(".pdf") else "application/octet-stream"
                url = f"data:{_mime};base64," + _b64mod.b64encode(plan_up.getvalue()).decode()
            if url:
                db.table("dateien").insert({
                    "projekt_id": pid, "kategorie": "Plaene",
                    "unterordner": plan_ord, "datei_name": plan_up.name, "datei_url": url
                }).execute()
                st.success("Plan gespeichert!")
                st.rerun()

    st.markdown("---")
    try:
        plaene = db.table("dateien").select("*").eq("projekt_id", pid).eq("kategorie", "Plaene").eq("unterordner", plan_ord).neq("datei_name", ".ordner").order("erstellt_am", desc=True).execute().data
        if plaene:
            for p in plaene:
                with st.container(border=True):
                    col_n, col_dl, col_del = st.columns([5, 2, 1])
                    with col_n:
                        st.markdown(f"**{p['datei_name']}**")
                        st.caption(p["erstellt_am"][:10] if p.get("erstellt_am") else "")
                    with col_dl:
                        if p["datei_url"].startswith("http"):
                            st.markdown(f"[Download]({p['datei_url']})")
                        elif p["datei_url"].startswith("data:"):
                            import base64 as _b64dl2
                            try:
                                _fb2 = _b64dl2.b64decode(p["datei_url"].split(",")[1])
                                st.download_button("⬇ Download", data=_fb2, file_name=p["datei_name"],
                                                   use_container_width=True, key=f"dl_plan_{p['id']}")
                            except Exception:
                                pass
                    with col_del:
                        if st.button("Lösch.", key=f"del_plan_{p['id']}", help="Löschen"):
                            db.table("dateien").delete().eq("id", p["id"]).execute()
                            st.rerun()
                    # PDF inline anzeigen
                    if _HAS_PDF and p["datei_name"].lower().endswith(".pdf"):
                        with st.expander("PDF anzeigen"):
                            try:
                                if p["datei_url"].startswith("data:"):
                                    import base64 as _b64
                                    pdf_bytes = _b64.b64decode(p["datei_url"].split(",")[1])
                                    _pdf_viewer(input=pdf_bytes, width=700)
                                elif p["datei_url"].startswith("http"):
                                    import requests as _req
                                    pdf_bytes = _req.get(p["datei_url"], timeout=8).content
                                    _pdf_viewer(input=pdf_bytes, width=700)
                            except Exception as ex:
                                st.caption(f"Vorschau nicht verfügbar: {ex}")
        else:
            st.markdown(f"""
            <div style='text-align:center; padding:2rem; color:#AAAAAA;'>
                <p style='font-size:0.9rem; margin:0;'>Noch keine Pläne in <b>'{plan_ord}'</b></p>
            </div>""", unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Fehler: {e}")

# ═════════════════════════════════════════════════════════════
# TAB 3: VERTRÄGE
# ═════════════════════════════════════════════════════════════
with tab_vertraege:
    st.markdown("### Verträge & Vereinbarungen")
    st.caption("Verträge, Leistungsverzeichnisse und Vereinbarungen für dieses Projekt.")

    # ── Briefkopf-Vorlage ─────────────────────────────────────
    with st.expander("🖼️  Briefkopf-Vorlage (für VOB-Dokumente & Tagesberichte)", expanded=False):
        try:
            _bk_rows = db.table("dateien").select("*").eq("projekt_id", pid).eq("kategorie", "Briefkopf").execute().data
        except Exception:
            _bk_rows = []

        if _bk_rows:
            _bk = _bk_rows[0]
            st.success(f"Aktive Vorlage: **{_bk['datei_name']}**")
            st.caption("Alle Word-Exporte (VOB-Schriftverkehr, Tagesbericht) verwenden diesen Briefkopf automatisch.")
            if st.button("Vorlage entfernen", key="del_briefkopf"):
                db.table("dateien").delete().eq("id", _bk["id"]).execute()
                st.rerun()
        else:
            st.info("Keine Vorlage hinterlegt — Word-Dokumente werden im Standard-Layout erstellt.")

        st.markdown("""
        <div style='font-size:0.82rem; color:#777672; margin-bottom:0.5rem;'>
        Tipp: Erstellen Sie in Word ein leeres Dokument mit Ihrem Firmenlogo, Briefkopf und Fußzeile.
        Inhalt weglassen — die App befüllt das Dokument automatisch.
        </div>""", unsafe_allow_html=True)

        with st.form("briefkopf_form", clear_on_submit=True):
            bk_up = st.file_uploader("Word-Vorlage hochladen (.docx)", type=["docx"],
                                      help="Word-Dokument mit Briefkopf, Logo und Adresse")
            if st.form_submit_button("Vorlage speichern", type="primary") and bk_up:
                try:
                    db.table("dateien").delete().eq("projekt_id", pid).eq("kategorie", "Briefkopf").execute()
                except Exception:
                    pass
                import base64 as _b64mod
                _bk_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                _bk_url = f"data:{_bk_mime};base64," + _b64mod.b64encode(bk_up.getvalue()).decode()
                db.table("dateien").insert({
                    "projekt_id": pid, "kategorie": "Briefkopf",
                    "unterordner": "", "datei_name": bk_up.name, "datei_url": _bk_url
                }).execute()
                st.success(f"Vorlage '{bk_up.name}' gespeichert!")
                st.rerun()

    st.markdown("---")
    with st.form("vertrag_form", clear_on_submit=True):
        vert_up = st.file_uploader("Dokument hochladen (PDF, Word)", type=["pdf","docx","doc"])
        if st.form_submit_button("Speichern", type="primary") and vert_up:
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            pfad = f"{pid}/Vertraege/{ts}_{vert_up.name}"
            if modus == "supabase":
                url = upload_datei(vert_up.getvalue(), pfad)
            else:
                import base64 as _b64mod
                _ext = vert_up.name.lower()
                if _ext.endswith(".pdf"):
                    _mime = "application/pdf"
                elif _ext.endswith(".docx"):
                    _mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                else:
                    _mime = "application/octet-stream"
                url = f"data:{_mime};base64," + _b64mod.b64encode(vert_up.getvalue()).decode()
            if url:
                db.table("dateien").insert({
                    "projekt_id": pid, "kategorie": "Vertraege",
                    "unterordner": "", "datei_name": vert_up.name, "datei_url": url
                }).execute()
                st.success("Dokument gespeichert!")
                st.rerun()

    st.markdown("---")
    try:
        vertraege = db.table("dateien").select("*").eq("projekt_id", pid).eq("kategorie", "Vertraege").neq("datei_name", ".ordner").order("erstellt_am", desc=True).execute().data
        if vertraege:
            for v in vertraege:
                with st.container(border=True):
                    col_n, col_dl, col_del = st.columns([5, 2, 1])
                    with col_n:
                        st.markdown(f"**{v['datei_name']}**")
                        st.caption(v["erstellt_am"][:10] if v.get("erstellt_am") else "")
                    with col_dl:
                        if v["datei_url"].startswith("http"):
                            st.markdown(f"[Download]({v['datei_url']})")
                        elif v["datei_url"].startswith("data:"):
                            import base64 as _b64dl
                            try:
                                _fb = _b64dl.b64decode(v["datei_url"].split(",")[1])
                                st.download_button("⬇ Download", data=_fb, file_name=v["datei_name"],
                                                   use_container_width=True, key=f"dl_vert_{v['id']}")
                            except Exception:
                                pass
                    with col_del:
                        if st.button("Lösch.", key=f"del_vert_{v['id']}", help="Löschen"):
                            db.table("dateien").delete().eq("id", v["id"]).execute()
                            st.rerun()
                    # PDF inline anzeigen
                    if _HAS_PDF and v["datei_name"].lower().endswith(".pdf"):
                        with st.expander("PDF anzeigen"):
                            try:
                                if v["datei_url"].startswith("data:"):
                                    import base64 as _b64
                                    pdf_bytes = _b64.b64decode(v["datei_url"].split(",")[1])
                                    _pdf_viewer(input=pdf_bytes, width=700)
                                elif v["datei_url"].startswith("http"):
                                    import requests as _req
                                    pdf_bytes = _req.get(v["datei_url"], timeout=8).content
                                    _pdf_viewer(input=pdf_bytes, width=700)
                            except Exception as ex:
                                st.caption(f"Vorschau nicht verfügbar: {ex}")
        else:
            st.markdown("""
            <div style='text-align:center; padding:2rem; color:#AAAAAA;'>
                <p style='font-size:0.9rem; margin:0;'>Noch keine Verträge hochgeladen</p>
            </div>""", unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Fehler: {e}")

# ═════════════════════════════════════════════════════════════
# TAB 4: KI-CHAT
# ═════════════════════════════════════════════════════════════
with tab_chat:
    st.markdown("### KI-Assistent")
    st.caption(f"Stelle Fragen zu Projekt '{pname}', Baurecht, Verträgen und Abläufen.")

    key_verlauf = f"chat_{pid}"
    if key_verlauf not in st.session_state:
        st.session_state[key_verlauf] = []

    def lade_kontext() -> str:
        try:
            alle = db.table("dateien").select("datei_name, kategorie, datei_url").eq("projekt_id", pid).neq("datei_name", ".ordner").execute().data
            if not alle:
                return "Keine Dokumente vorhanden."
            gruppen = {}
            for d in alle:
                gruppen.setdefault(d["kategorie"], []).append(d["datei_name"])
            zeilen = []
            for kat, namen in gruppen.items():
                zeilen.append(f"{kat}: {', '.join(namen)}")
            kontext = "Dokumente im Projekt:\n" + "\n".join(zeilen)
            # Vertragstext extrahieren und der KI übergeben
            # (bis zu 6 Dokumente, je bis 20.000 Zeichen — Gemini 2.5 Flash hat sehr großen Kontext)
            for v in [d for d in alle if d.get("kategorie") == "Vertraege" and d.get("datei_url")][:6]:
                try:
                    import base64 as _b64ctx, io as _io_ctx
                    url = v["datei_url"]
                    if url.startswith("data:"):
                        raw = _b64ctx.b64decode(url.split(",")[1])
                    elif url.startswith("http"):
                        import requests as _rq
                        raw = _rq.get(url, timeout=5).content
                    else:
                        continue
                    name = v["datei_name"].lower()
                    vtext = ""
                    if name.endswith(".docx"):
                        from docx import Document as _DocCtx
                        vtext = "\n".join(p.text for p in _DocCtx(_io_ctx.BytesIO(raw)).paragraphs if p.text.strip())
                    elif name.endswith(".txt"):
                        vtext = raw.decode("utf-8", errors="replace")
                    elif name.endswith(".pdf"):
                        try:
                            from pdfminer.high_level import extract_text as _pdf_ex
                            vtext = _pdf_ex(_io_ctx.BytesIO(raw))
                        except ImportError:
                            pass
                    if vtext:
                        kontext += f"\n\n--- VERTRAGSINHALT: {v['datei_name']} ---\n{vtext[:20000]}"
                except Exception:
                    pass
            return kontext
        except Exception:
            return ""

    for msg in st.session_state[key_verlauf]:
        with st.chat_message(msg["rolle"]):
            st.markdown(msg["text"])

    frage = st.chat_input("Frage stellen — z.B. 'Schulde ich das Baumfällen?' oder 'Was sind typische Nachtragsgründe?'")
    if frage:
        st.session_state[key_verlauf].append({"rolle": "user", "text": frage})
        with st.chat_message("user"):
            st.markdown(frage)
        with st.chat_message("assistant"):
            with st.spinner("Denke nach..."):
                import time as _time
                # Rate-Limiting: min. 4 Sek. zwischen Anfragen
                rate_key = f"last_ki_call_{pid}"
                last_call = st.session_state.get(rate_key, 0)
                elapsed = _time.time() - last_call
                if elapsed < 4:
                    wartezeit = int(4 - elapsed) + 1
                    st.warning(f"Bitte {wartezeit} Sekunde(n) warten bevor du die nächste Frage sendest.")
                else:
                    try:
                        kontext = lade_kontext()
                        verlauf = "\n".join([f"{m['rolle'].upper()}: {m['text']}"
                                             for m in st.session_state[key_verlauf][-6:]])
                        prompt = f"""Du bist Baurechts-Assistent für Projekt '{pname}'.
Kenntnisse: VOB/B (§§ 1-18), BGB Werkvertrag (§§ 631-650v), Baupraxis.
Wichtig: §2 Abs.5/6 VOB/B Nachtrag VOR Ausführung! §6 Abs.1 Behinderung sofort anzeigen! §12 Abnahme = Fälligkeit. §13 Mängel 4 Jahre. §16 Zahlung.
Projektdokumente: {kontext}
Gespräch: {verlauf}
Antworte auf Deutsch, kurz, §-Angabe bei Rechtsfragen, JA/NEIN + Begründung."""
                        antwort = _ki_call(prompt)
                        st.session_state[rate_key] = _time.time()
                        st.markdown(antwort)
                        st.session_state[key_verlauf].append({"rolle": "assistant", "text": antwort})
                    except Exception as e:
                        err = str(e)
                        if "GEMINI_API_KEY" in err or "401" in err or "invalid_api_key" in err.lower() or "API_KEY_INVALID" in err:
                            st.error("**API-Key ungültig** — Bitte GEMINI_API_KEY in den Streamlit Secrets prüfen.")
                        elif "429" in err or "rate_limit" in err.lower():
                            st.session_state[rate_key] = _time.time()
                            st.warning(
                                "**KI-Limit erreicht** — Bitte **60 Sekunden warten** und erneut versuchen.  \n"
                                "Tipp: Stelle mehrere Fragen in einer Nachricht, um Anfragen zu sparen."
                            )
                        else:
                            st.error(f"KI-Fehler: {err[:300]}")

    col_clear, _ = st.columns([2, 8])
    with col_clear:
        if st.button("Chat leeren", key="clear_chat"):
            st.session_state[key_verlauf] = []
            st.rerun()

# ═════════════════════════════════════════════════════════════
# TAB 5: TO-DO
# ═════════════════════════════════════════════════════════════
with tab_todos:
    st.markdown("### To-Do Liste")

    STANDARD_TODOS = {
        "Arbeitsvorbereitung": [
            "Baustelleneinrichtungsplan erstellen",
            "Subunternehmer beauftragen",
            "Pläne freigeben lassen",
            "Materialbestellung vorbereiten",
            "Einweisungen und Sicherheitsunterweisung planen",
        ],
        "Bauausführung": [
            "Tagesbericht führen",
            "Qualitätskontrolle durchführen",
            "Mängel dokumentieren",
            "Abrechnung der Subunternehmer prüfen",
            "Besprechungsprotokolle führen",
        ],
        "Abnahme & Übergabe": [
            "Begehung mit Auftraggeber planen",
            "Restmängelliste erstellen",
            "Dokumentation zusammenstellen",
            "Schlussrechnung stellen",
            "Bürgschaften zurückfordern",
        ],
    }

    try:
        vorhanden = db.table("todos").select("id").eq("projekt_id", pid).execute().data
        if not vorhanden:
            for phase, aufgaben in STANDARD_TODOS.items():
                for a in aufgaben:
                    db.table("todos").insert({"projekt_id": pid, "titel": a, "phase": phase}).execute()
    except Exception:
        pass

    with st.expander("Eigenes To-Do hinzufügen"):
        with st.form("todo_neu", clear_on_submit=True):
            todo_titel = st.text_input("Aufgabe *")
            todo_phase = st.selectbox("Phase", list(STANDARD_TODOS.keys()) + ["Allgemein"])
            todo_beschr = st.text_area("Beschreibung (optional)", height=70)
            if st.form_submit_button("Hinzufügen", type="primary") and todo_titel.strip():
                db.table("todos").insert({
                    "projekt_id": pid, "titel": todo_titel.strip(),
                    "phase": todo_phase, "beschreibung": todo_beschr
                }).execute()
                st.rerun()

    try:
        alle_todos = db.table("todos").select("*").eq("projekt_id", pid).order("erstellt_am").execute().data
        if alle_todos:
            ges = len(alle_todos)
            erl = sum(1 for t in alle_todos if t["erledigt"])
            pct = int(erl / ges * 100)
            st.markdown(f"""
            <div style='background:white; border:1px solid #e8edf3; border-radius:12px;
                        padding:1rem 1.2rem; margin-bottom:1rem; display:flex;
                        align-items:center; gap:1rem;'>
                <div style='flex:1;'>
                    <div style='font-size:0.8rem; color:#64748b; font-weight:500; margin-bottom:0.4rem;'>
                        Gesamtfortschritt</div>
                    <div style='background:#f1f5f9; border-radius:999px; height:8px; overflow:hidden;'>
                        <div style='background:linear-gradient(90deg,#F07030,#F59050);
                             width:{pct}%; height:100%; border-radius:999px;
                             transition:width 0.3s;'></div>
                    </div>
                </div>
                <div style='font-size:1.3rem; font-weight:800; color:#1e293b; white-space:nowrap;'>
                    {erl}/{ges} <span style='font-size:0.9rem; color:#64748b;'>({pct}%)</span>
                </div>
            </div>
            """, unsafe_allow_html=True)

            phasen = list(dict.fromkeys([t["phase"] for t in alle_todos]))
            for phase in phasen:
                phase_todos = [t for t in alle_todos if t["phase"] == phase]
                erl_n = sum(1 for t in phase_todos if t["erledigt"])
                with st.expander(f"{phase} — {erl_n}/{len(phase_todos)} erledigt", expanded=True):
                    for t in phase_todos:
                        col_check, col_text, col_del = st.columns([1, 9, 1])
                        with col_check:
                            neu = st.checkbox("Erledigt", value=t["erledigt"], key=f"todo_check_{t['id']}", label_visibility="collapsed")
                            if neu != t["erledigt"]:
                                db.table("todos").update({"erledigt": neu}).eq("id", t["id"]).execute()
                                st.rerun()
                        with col_text:
                            titel_safe = html_mod.escape(t["titel"])
                            if t["erledigt"]:
                                st.markdown(f"<span style='text-decoration:line-through; color:#AAAAAA;'>{titel_safe}</span>", unsafe_allow_html=True)
                            else:
                                st.markdown(f"<span style='color:#1e293b;'>{titel_safe}</span>", unsafe_allow_html=True)
                            if t.get("beschreibung"):
                                st.caption(t["beschreibung"])
                        with col_del:
                            if st.button("X", key=f"del_todo_{t['id']}", help="Löschen"):
                                db.table("todos").delete().eq("id", t["id"]).execute()
                                st.rerun()
    except Exception as e:
        st.error(f"Fehler: {e}")

# ═════════════════════════════════════════════════════════════
# TAB 6: KALENDER
# ═════════════════════════════════════════════════════════════
with tab_kalender:
    st.markdown("### Baustellen-Kalender")

    FARBEN = {
        "gruen":  ("Info",             "#dcfce7", "#166534", "#16a34a"),
        "blau":   ("Bauleiter-Termin", "#dbeafe", "#1e3a8a", "#2563eb"),
        "orange": ("Polier-Termin",    "#fff7ed", "#7c2d12", "#ea580c"),
    }

    col_leg1, col_leg2, col_leg3 = st.columns(3)
    for col, (k, (label, bg, fg, acc_)) in zip([col_leg1, col_leg2, col_leg3], FARBEN.items()):
        with col:
            st.markdown(
                f"<div style='background:{bg};color:{fg};padding:8px 14px;border-radius:8px;"
                f"font-size:0.85rem;font-weight:600;text-align:center;'>{label}</div>",
                unsafe_allow_html=True
            )
    st.markdown("")

    with st.expander("Termin eintragen"):
        with st.form("termin_neu", clear_on_submit=True):
            t_titel = st.text_input("Bezeichnung *", placeholder="z.B. Betonlieferung, Jour Fixe, Behinderung Rohbau")
            col_d1, col_d2 = st.columns(2)
            with col_d1:
                t_datum     = st.date_input("Beginn *", value=datetime.date.today(), format="DD.MM.YYYY")
            with col_d2:
                t_datum_bis = st.date_input("Ende (leer = eintägig)", value=None, format="DD.MM.YYYY")
            col_tv, col_tb = st.columns(2)
            with col_tv:
                t_von = st.text_input("Uhrzeit Von", placeholder="08:00")
            with col_tb:
                t_bis = st.text_input("Uhrzeit Bis", placeholder="16:00")
            t_kat = st.radio("Kategorie", ["gruen", "blau", "orange"],
                             format_func=lambda k: FARBEN[k][0],
                             horizontal=True)
            t_beschr = st.text_area("Notizen", height=60)
            if st.form_submit_button("Eintragen", type="primary") and t_titel.strip():
                db.table("kalender").insert({
                    "projekt_id": pid, "titel": t_titel.strip(),
                    "datum": str(t_datum),
                    "datum_bis": str(t_datum_bis) if t_datum_bis else None,
                    "uhrzeit_von": t_von, "uhrzeit_bis": t_bis,
                    "kategorie": t_kat, "beschreibung": t_beschr
                }).execute()
                st.success("Termin eingetragen!")
                st.rerun()

    try:
        alle_termine = db.table("kalender").select("*").eq("projekt_id", pid).order("datum").execute().data
        farb_map = {"gruen": "#16a34a", "blau": "#2563eb", "orange": "#ea580c"}

        if _HAS_CALENDAR:
            events = []
            for t in alle_termine:
                end = t.get("datum_bis") or t["datum"]
                # FullCalendar: end-Datum bei mehrtägigen Events ist exklusiv → +1 Tag
                try:
                    dt_end = datetime.date.fromisoformat(end)
                    dt_end += datetime.timedelta(days=1)
                    end_str = str(dt_end)
                except Exception:
                    end_str = end
                events.append({
                    "title": t["titel"],
                    "start": t["datum"],
                    "end": end_str,
                    "color": farb_map.get(t["kategorie"], "#2563eb"),
                    "extendedProps": {"beschreibung": t.get("beschreibung", ""), "db_id": t["id"]},
                })
            cal_opts = {
                "headerToolbar": {"left": "prev,next today", "center": "title", "right": "dayGridMonth,listMonth"},
                "initialView": "dayGridMonth",
                "locale": "de",
                "height": 560,
                "editable": False,
                "selectable": False,
            }
            cal_result = st_calendar(events=events, options=cal_opts, key=f"cal_{pid}")
            # Geklickten Termin anzeigen
            if cal_result and cal_result.get("eventClick"):
                ev = cal_result["eventClick"].get("event", {})
                props = ev.get("extendedProps", {})
                with st.container(border=True):
                    st.markdown(f"**{ev.get('title', '')}**")
                    st.caption(ev.get("start", ""))
                    if props.get("beschreibung"):
                        st.markdown(props["beschreibung"])
                    if st.button("Termin löschen", key="del_kal_click"):
                        try:
                            db.table("kalender").delete().eq("id", props["db_id"]).execute()
                            st.rerun()
                        except Exception:
                            pass
        else:
            # Fallback: Liste
            if not alle_termine:
                st.info("Keine Termine eingetragen.")
            for t in alle_termine:
                bg = FARBEN[t["kategorie"]][1]; fg = FARBEN[t["kategorie"]][2]; acc = FARBEN[t["kategorie"]][3]
                col_t, col_del = st.columns([12, 1])
                with col_t:
                    st.markdown(f"<div style='background:{bg};color:{fg};padding:8px 14px;border-radius:8px;border-left:3px solid {acc};margin:2px 0'><b>{html_mod.escape(t['titel'])}</b> <span style='font-size:0.82rem'>{t['datum']}</span></div>", unsafe_allow_html=True)
                with col_del:
                    if st.button("X", key=f"del_kal_{t['id']}"):
                        db.table("kalender").delete().eq("id", t["id"]).execute()
                        st.rerun()
    except Exception as e:
        st.error(f"Fehler: {e}")

# ═════════════════════════════════════════════════════════════
# TAB 7: VOB-SCHRIFTVERKEHR
# ═════════════════════════════════════════════════════════════
with tab_vob:
    st.markdown("### VOB-Schriftverkehr")
    st.caption("KI-gestützte Erstellung von Mehrkostenanzeige und Behinderungsanzeige nach VOB/B.")

    _vob_bauleiter = get_name() or pname

    # Briefkopf-Status anzeigen
    try:
        _hat_vorlage = bool(db.table("dateien").select("id").eq("projekt_id", pid).eq("kategorie", "Briefkopf").execute().data)
    except Exception:
        _hat_vorlage = False
    if _hat_vorlage:
        st.markdown("""<div style='background:#f0fdf4;border:1px solid #bbf7d0;border-radius:8px;
            padding:8px 14px;font-size:0.83rem;color:#166534;margin-bottom:0.8rem;'>
            ✓ Briefkopf-Vorlage aktiv — Word-Exporte verwenden automatisch Ihren Briefkopf
            </div>""", unsafe_allow_html=True)

    vob_typ = st.radio(
        "Dokument erstellen",
        ["Mehrkostenanzeige", "Behinderungsanzeige"],
        horizontal=True,
        key="vob_typ"
    )

    # ── Briefkopf-Vorlage laden ───────────────────────────────
    def _lade_briefkopf() -> bytes:
        try:
            rows = db.table("dateien").select("datei_url").eq("projekt_id", pid).eq("kategorie", "Briefkopf").execute().data
            if not rows:
                return None
            url = rows[0]["datei_url"]
            if url.startswith("data:"):
                import base64 as _b64
                return _b64.b64decode(url.split(",")[1])
            elif url.startswith("http"):
                import requests as _req
                return _req.get(url, timeout=8).content
        except Exception:
            return None
        return None

    # ── Hilfsfunktion: Word-Dokument (professionelles Brief-Layout) ──
    def _erstelle_docx(titel_dok: str, meta: dict, text: str, bilder: list = None) -> bytes:
        from docx import Document
        from docx.shared import Cm, Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement
        import io, re as _re_md

        ACCENT = RGBColor(0xF0, 0x70, 0x30)   # R·APP-Orange für Akzentlinien
        DARK   = RGBColor(0x1A, 0x1A, 0x1A)   # Anthrazit für Titel/Überschriften

        vorlage_bytes = _lade_briefkopf()
        hat_vorlage = bool(vorlage_bytes)
        if vorlage_bytes:
            doc = Document(io.BytesIO(vorlage_bytes))
            body = doc.element.body
            for child in list(body):
                if child.tag != _qn('w:sectPr'):
                    body.remove(child)
        else:
            doc = Document()
            s = doc.sections[0]
            s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
            s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.0)

        # Basis-Typografie: konsistente Schrift + angenehmer Zeilenabstand.
        # Schriftfamilie nur setzen, wenn KEINE Kunden-Vorlage (deren Stil bleibt erhalten).
        try:
            normal = doc.styles['Normal']
            if not hat_vorlage:
                normal.font.name = 'Calibri'
            normal.font.size = Pt(11)
            normal.paragraph_format.space_before = Pt(0)
            normal.paragraph_format.space_after = Pt(6)
            normal.paragraph_format.line_spacing = 1.18
        except Exception:
            pass

        def _p(txt="", *, bold=False, italic=False, size=None, align=None,
               space_after=6, space_before=0, color=None):
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(space_before)
            para.paragraph_format.space_after = Pt(space_after)
            if align is not None:
                para.alignment = align
            if txt:
                r = para.add_run(txt)
                r.bold = bold; r.italic = italic
                if size:  r.font.size = Pt(size)
                if color: r.font.color.rgb = color
            return para

        # Dünne horizontale Linie (Absatz-Unterrand) statt Strich-Zeichen
        def _hr(color="BFBFBF", sz="6", space_after=8, space_before=2):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            pPr = p._p.get_or_add_pPr()
            pbdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(_qn('w:val'), 'single'); bottom.set(_qn('w:sz'), sz)
            bottom.set(_qn('w:space'), '1');    bottom.set(_qn('w:color'), color)
            pbdr.append(bottom); pPr.append(pbdr)
            return p

        # ── Titel + Akzentlinie ───────────────────────────────
        _p(titel_dok, bold=True, size=17, color=DARK,
           align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
        _hr(color="F07030", sz="18", space_after=8, space_before=0)   # kräftige Orange-Linie

        # ── Metadaten-Block: 2-spaltig (kompakt, Infos links/rechts) ──
        meta_items = [(k, str(v)) for k, v in (meta or {}).items() if v]
        if meta_items:
            import math as _math
            rows_n = _math.ceil(len(meta_items) / 2)
            mt = doc.add_table(rows=rows_n, cols=2)
            mt.autofit = True
            for idx, (k, v) in enumerate(meta_items):
                r_i = idx % rows_n      # spaltenweise: erst linke Spalte füllen, dann rechte
                c_i = idx // rows_n
                para = mt.rows[r_i].cells[c_i].paragraphs[0]
                para.paragraph_format.space_after = Pt(1)
                rk = para.add_run(f"{k}:  "); rk.bold = True; rk.font.size = Pt(10); rk.font.color.rgb = DARK
                rv = para.add_run(v);          rv.font.size = Pt(10)
            _hr(color="D9D9D9", sz="6", space_after=12, space_before=6)

        # ── Fließtext im Blocksatz, mit Überschriften-Hierarchie ──
        text = _re_md.sub(r'\n{3,}', '\n\n', text.strip())
        for zeile in text.split("\n"):
            z = zeile.strip()
            if not z:
                continue
            if z.startswith("## "):
                _p(z[3:], bold=True, size=12, color=DARK, space_before=10, space_after=3)
            elif z.startswith("# "):
                _p(z[2:], bold=True, size=13, color=DARK, space_before=12, space_after=4)
            elif "**" in z:
                p2 = doc.add_paragraph()
                p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p2.paragraph_format.space_after = Pt(6)
                for _i, _part in enumerate(_re_md.split(r'\*\*(.+?)\*\*', z)):
                    if _part:
                        p2.add_run(_part).bold = (_i % 2 == 1)
            else:
                _p(z, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6)

        # ── Unterschriftenzeile (zweispaltig, sauber) ─────────
        _hr(color="D9D9D9", sz="6", space_after=20, space_before=14)
        sig = doc.add_table(rows=2, cols=2)
        sig.autofit = True
        sig.rows[0].cells[0].text = "_______________________________"
        sig.rows[0].cells[1].text = "_______________________________"
        c10 = sig.rows[1].cells[0].paragraphs[0]; r10 = c10.add_run("Ort, Datum"); r10.font.size = Pt(9); r10.font.color.rgb = RGBColor(0x80,0x80,0x80)
        c11 = sig.rows[1].cells[1].paragraphs[0]; r11 = c11.add_run("Unterschrift Bauleiter"); r11.font.size = Pt(9); r11.font.color.rgb = RGBColor(0x80,0x80,0x80)

        # ── Foto-Anhang ───────────────────────────────────────
        if bilder:
            doc.add_page_break()
            _p("Foto-Anhang", bold=True, size=13, color=DARK, space_after=8)
            _hr(color="D9D9D9", sz="6", space_after=12)
            for i, bild_bytes in enumerate(bilder, 1):
                try:
                    _p(f"Foto {i}", bold=True, size=10, space_after=3, space_before=8)
                    doc.add_picture(io.BytesIO(bild_bytes), width=Inches(5.8))
                except Exception:
                    pass

        buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

    def _ki_fehler(e):
        err = str(e)
        if "GEMINI_API_KEY" in err or "401" in err or "invalid_api_key" in err.lower() or "API_KEY_INVALID" in err:
            st.error("**API-Key ungültig** — Bitte GEMINI_API_KEY in den Streamlit Secrets prüfen.")
        elif "429" in err or "rate_limit" in err.lower():
            st.warning("KI-Limit erreicht — Bitte 1 Minute warten und erneut versuchen.")
        else:
            st.error(f"KI-Fehler: {err[:300]}")

    def _ki_generieren(prompt: str) -> str:
        return _ki_call(prompt)

    def _ergebnis_anzeigen(key: str, titel_dok: str, meta: dict):
        result = st.session_state.get(key)
        if not result:
            return
        titel_anzeige = result.get("titel", titel_dok)
        st.markdown("---")
        st.markdown(f"#### Neu generiert: {titel_anzeige}")
        st.markdown(result["text"])
        st.markdown("---")
        c1, c2, c3 = st.columns([3, 3, 1])
        with c1:
            fname = titel_anzeige.replace(" ", "_")
            st.download_button("Als .txt", data=result["text"],
                file_name=f"{fname}_{result.get('datum','')}.txt", mime="text/plain",
                use_container_width=True, key=f"neu_txt_{key}")
        with c2:
            try:
                docx_b = _erstelle_docx(titel_anzeige, meta, result["text"],
                                        bilder=result.get("bilder") or [])
                st.download_button("Als Word (.docx)", data=docx_b,
                    file_name=f"{fname}_{result.get('datum','')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key=f"neu_docx_{key}")
            except Exception as ex:
                st.error(f"Word-Fehler: {ex}")
        with c3:
            if st.button("Schliessen", key=f"clear_{key}", use_container_width=True):
                del st.session_state[key]
                st.rerun()

    def _zeige_archiv(typ: str, anzeige_name: str):
        docs = lade_dokumente(typ)
        if not docs:
            return
        st.markdown("---")
        st.markdown(f"#### Archiv — {anzeige_name}")
        for doc in reversed(docs):
            meta_d = {}
            try:
                meta_d = _json.loads(doc.get("meta", "{}"))
            except Exception:
                pass
            datum_str = meta_d.get("datum", doc.get("erstellt_am", "")[:10])
            label = f"{doc['titel']} — {datum_str}"
            with st.expander(label):
                st.markdown(doc["inhalt"])
                st.markdown("---")
                c1, c2, c3 = st.columns([3, 3, 1])
                fname = doc["titel"].replace(" ", "_")
                with c1:
                    st.download_button("Als .txt", data=doc["inhalt"],
                        file_name=f"{fname}.txt", mime="text/plain",
                        use_container_width=True, key=f"dl_txt_{doc['id']}")
                with c2:
                    try:
                        meta_dok = {
                            "Projekt": pname,
                            "Datum": datum_str,
                            "Auftraggeber": meta_d.get("auftraggeber", ""),
                            "Bauleiter": _vob_bauleiter,
                        }
                        docx_b = _erstelle_docx(doc["titel"], meta_dok, doc["inhalt"])
                        st.download_button("Als Word (.docx)", data=docx_b,
                            file_name=f"{fname}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True, key=f"dl_docx_{doc['id']}")
                    except Exception as ex:
                        st.error(f"Word-Fehler: {ex}")
                with c3:
                    if st.button("Löschen", key=f"del_dok_{doc['id']}", use_container_width=True):
                        try:
                            db.table("schriftverkehr").delete().eq("id", doc["id"]).execute()
                        except Exception:
                            pass
                        st.rerun()

    # ══════════════════════════════════════════════════════════
    # MEHRKOSTENANZEIGE (§ 2 Abs. 5 / 6 VOB/B)
    # ══════════════════════════════════════════════════════════
    if vob_typ == "Mehrkostenanzeige":
        mk_grundlage = st.radio("Rechtsgrundlage", ["§ 2 Abs. 5 VOB/B", "§ 2 Abs. 6 VOB/B", "§ 650b BGB"],
                                horizontal=True, key="mk_grundlage")
        _mk_info = {
            "§ 2 Abs. 5 VOB/B": "Geänderte Leistung (z. B. abweichender Boden, geänderte Ausführung).",
            "§ 2 Abs. 6 VOB/B": "Zusätzliche, im Vertrag nicht vereinbarte Leistung.",
            "§ 650b BGB": "Änderungsanordnung im BGB-Bauvertrag; Vergütungsanpassung nach § 650c BGB.",
        }.get(mk_grundlage, "")
        st.markdown(f"""
        <div style='background:#F7F6F4;border:1px solid #E8E6E2;border-left:3px solid #F07030;
                    border-radius:8px;padding:0.9rem 1.1rem;margin:0.6rem 0 1rem 0;
                    font-size:0.84rem;color:#444441;'>
        <b style="color:#1A1A1A;">Gewählte Grundlage: {mk_grundlage}</b> — {_mk_info}<br>
        <b style="color:#F07030;">Wichtig:</b> Die Anzeige muss <u>vor Ausführung</u> erstattet werden. Ohne Ankündigung entfällt der Vergütungsanspruch.
        </div>
        """, unsafe_allow_html=True)

        with st.container(border=True):
            c1, c2 = st.columns(2)
            with c1:
                mk_datum = st.date_input("Datum", value=datetime.date.today(), format="DD.MM.YYYY", key="mk_datum")
            with c2:
                mk_ag = st.text_input("Auftraggeber", value=projekt.get("auftraggeber",""), placeholder="z.B. Stadtwerke Konstanz GmbH", key="mk_ag")
            mk_vertrag = st.text_input("Vertrags-/Auftragsnummer", value=projekt.get("vertragsnummer",""), placeholder="z.B. V-2026-042", key="mk_vertrag")

            mk_beschr = st.text_area(
                "Beschreibung der Mehrleistung *",
                placeholder="z.B. Aufgrund nicht in der Ausschreibung enthaltener Felsschichten (Bodenklasse 7) "
                            "müssen statt der geplanten 3 Handschachtungen insgesamt 8 Baggerschachtungen durchgeführt werden. "
                            "Die Abweichung wurde am 03.06.2026 festgestellt.",
                height=130, key="mk_beschr"
            )
            mk_betrag = st.text_input("Voraussichtliche Mehrkosten (ca.)",
                placeholder="z.B. ca. 8.500,– EUR netto", key="mk_betrag")
            mk_fotos = st.file_uploader(
                "Foto-Anhang (optional)", type=["jpg","jpeg","png","webp"],
                accept_multiple_files=True, key="mk_fotos",
                help="Fotos werden als Anhang ins Word-Dokument eingefügt"
            )

            if st.button("Mehrkostenanzeige generieren", type="primary",
                         disabled=not mk_beschr.strip(), use_container_width=True, key="mk_btn"):
                with st.spinner("Erstelle Mehrkostenanzeige nach VOB/B..."):
                    try:
                        prompt = f"""Erstelle eine rechtssichere Mehrkostenanzeige nach {mk_grundlage} (VOB/B).

Auftragnehmer: {_vob_bauleiter}, Bauleitung {pname}
Auftraggeber: {mk_ag or 'Auftraggeber'}
Datum: {mk_datum.strftime('%d.%m.%Y')}
Auftrag: {mk_vertrag or '–'}
Sachverhalt: {mk_beschr}
Voraussichtliche Mehrkosten: {mk_betrag or 'werden noch ermittelt'}

Struktur: Briefkopf → Betreff "Mehrkostenanzeige gem. {mk_grundlage} – {pname}" → Anrede → Sachverhalt (was vereinbart / was abweicht) → Rechtsgrundlage → Mehrkosten-Ankündigung mit Vorbehalt → Aufforderung schriftliche Beauftragung VOR Ausführung (Frist 5 Werktage) → Grußformel.

Ton: sachlich-bestimmt, kein Konjunktiv. "Hiermit zeigen wir an…", "Wir behalten uns vor…"
Format: vollständiger Geschäftsbrief, kein Markdown."""
                        antwort = _ki_generieren(prompt)
                        nr = get_naechste_nummer("mehrkostenanzeige")
                        titel_dok = f"Mehrkostenanzeige Nr. {nr}"
                        speichere_dokument("mehrkostenanzeige", nr, titel_dok, antwort,
                                           {"datum": str(mk_datum), "auftraggeber": mk_ag,
                                            "grundlage": mk_grundlage, "betrag": mk_betrag})
                        st.session_state[f"mk_{pid}"] = {
                            "text": antwort, "datum": str(mk_datum),
                            "projekt": pname, "auftraggeber": mk_ag,
                            "titel": titel_dok,
                            "bilder": [f.getvalue() for f in mk_fotos] if mk_fotos else [],
                        }
                        st.rerun()
                    except Exception as e:
                        _ki_fehler(e)

        _ergebnis_anzeigen(f"mk_{pid}", "Mehrkostenanzeige",
                           {"Projekt": pname, "Bauleiter": _vob_bauleiter, "Auftraggeber": mk_ag, "Datum": str(mk_datum)})

        # Archiv Mehrkostenanzeigen
        _zeige_archiv("mehrkostenanzeige", "Mehrkostenanzeige")

    # ══════════════════════════════════════════════════════════
    # BEHINDERUNGSANZEIGE (§ 6 Abs. 1 VOB/B)
    # ══════════════════════════════════════════════════════════
    else:
        st.markdown("""
        <div style='background:#F7F6F4;border:1px solid #E8E6E2;border-left:3px solid #F07030;
                    border-radius:8px;padding:0.9rem 1.1rem;margin-bottom:1rem;
                    font-size:0.84rem;color:#444441;'>
        <b style="color:#1A1A1A;">Rechtliche Grundlage:</b> § 6 Abs. 1 VOB/B.<br>
        <b style="color:#F07030;">Wichtig:</b> Die Anzeige muss <u>unverzüglich</u> nach Eintritt der Behinderung erstattet werden. Nur dann bleibt der Anspruch auf Bauzeitverlängerung (§ 6 Abs. 2) und Schadensersatz (§ 6 Abs. 6) erhalten. Verspätete Anzeigen verwirken den Anspruch.
        </div>
        """, unsafe_allow_html=True)

        with st.container(border=True):
            c1, c2 = st.columns(2)
            with c1:
                bh_datum = st.date_input("Datum der Anzeige", value=datetime.date.today(), format="DD.MM.YYYY", key="bh_datum")
            with c2:
                bh_ag = st.text_input("Auftraggeber", value=projekt.get("auftraggeber",""), placeholder="z.B. Stadtwerke Konstanz GmbH", key="bh_ag")

            c3, c4 = st.columns(2)
            with c3:
                bh_beginn = st.date_input("Beginn der Behinderung", value=datetime.date.today(), format="DD.MM.YYYY", key="bh_beginn")
            with c4:
                bh_ursache_typ = st.selectbox(
                    "Art der Behinderungsursache",
                    ["Fehlendes Material/Ausführungsunterlagen vom AG",
                     "Witterungsbedingungen (Frost, Starkregen, etc.)",
                     "Andere Gewerke / Vorarbeiten nicht abgeschlossen",
                     "Baugenehmigung / Behördliche Auflagen",
                     "Unvorhersehbare Bodenverhältnisse",
                     "Anordnung des AG / Baustop",
                     "Sonstige"],
                    key="bh_ursache_typ"
                )

            bh_beschr = st.text_area(
                "Genaue Beschreibung der Behinderung *",
                placeholder="z.B. Am 03.06.2026 wurde festgestellt, dass die vom AG zu liefernden "
                            "Ausführungspläne für das Erdgeschoss bis heute nicht eingegangen sind. "
                            "Die Ausführung der Mauerarbeiten im EG ist dadurch vollständig blockiert. "
                            "Betroffen sind 4 Facharbeiter, die nicht anderweitig eingesetzt werden können.",
                height=130, key="bh_beschr"
            )

            c5, c6 = st.columns(2)
            with c5:
                bh_gewerke = st.text_input("Betroffene Gewerke/Leistungen",
                    placeholder="z.B. Mauerarbeiten EG, Bewehrungsarbeiten Decke", key="bh_gewerke")
            with c6:
                bh_dauer = st.text_input("Voraussichtliche Dauer",
                    placeholder="z.B. mindestens 5 Werktage", key="bh_dauer")

            bh_vertrag = st.text_input("Vertrags-/Auftragsnummer (optional)",
                value=projekt.get("vertragsnummer",""), placeholder="z.B. V-2026-042", key="bh_vertrag")
            bh_fotos = st.file_uploader(
                "Foto-Anhang (optional)", type=["jpg","jpeg","png","webp"],
                accept_multiple_files=True, key="bh_fotos",
                help="Fotos werden als Anhang ins Word-Dokument eingefügt"
            )

            if st.button("Behinderungsanzeige generieren", type="primary",
                         disabled=not bh_beschr.strip(), use_container_width=True, key="bh_btn"):
                with st.spinner("Erstelle Behinderungsanzeige nach § 6 Abs. 1 VOB/B..."):
                    try:
                        prompt = f"""Erstelle eine rechtssichere Behinderungsanzeige gem. § 6 Abs. 1 VOB/B.

Auftragnehmer: {_vob_bauleiter}, Bauleitung {pname}
Auftraggeber: {bh_ag or 'Auftraggeber'}
Datum der Anzeige: {bh_datum.strftime('%d.%m.%Y')}
Beginn der Behinderung: {bh_beginn.strftime('%d.%m.%Y')}
Ursache: {bh_ursache_typ}
Beschreibung: {bh_beschr}
Betroffene Leistungen: {bh_gewerke or 'laufende Arbeiten'}
Voraussichtliche Dauer: {bh_dauer or 'derzeit nicht abschätzbar'}
Auftrag: {bh_vertrag or '–'}

Struktur: Briefkopf → Betreff "Behinderungsanzeige gem. § 6 Abs. 1 VOB/B – {pname}" → Anrede → Anzeige der Behinderung → Sachverhalt → Auswirkungen auf Bauablauf → Forderungen: Fristverlängerung (§ 6 Abs. 2), Mehrkostenvorbehalt (§ 6 Abs. 6), Abhilfefrist 5 Werktage → Grußformel.

Ton: sachlich-bestimmt. "Hiermit zeigen wir an…", "Diese Anzeige erfolgt unverzüglich und fristwahrend."
Format: vollständiger Geschäftsbrief, kein Markdown."""
                        antwort = _ki_generieren(prompt)
                        nr = get_naechste_nummer("behinderungsanzeige")
                        titel_dok = f"Behinderungsanzeige Nr. {nr}"
                        speichere_dokument("behinderungsanzeige", nr, titel_dok, antwort,
                                           {"datum": str(bh_datum), "auftraggeber": bh_ag,
                                            "beginn": str(bh_beginn), "ursache": bh_ursache_typ})
                        st.session_state[f"bh_{pid}"] = {
                            "text": antwort, "datum": str(bh_datum),
                            "projekt": pname, "auftraggeber": bh_ag,
                            "titel": titel_dok,
                            "bilder": [f.getvalue() for f in bh_fotos] if bh_fotos else [],
                        }
                        st.rerun()
                    except Exception as e:
                        _ki_fehler(e)

        _ergebnis_anzeigen(f"bh_{pid}", "Behinderungsanzeige",
                           {"Projekt": pname, "Bauleiter": _vob_bauleiter, "Auftraggeber": bh_ag, "Datum": str(bh_datum)})

        # Archiv Behinderungsanzeigen
        _zeige_archiv("behinderungsanzeige", "Behinderungsanzeige")

# ═════════════════════════════════════════════════════════════
# TAB 8: TAGESBERICHT
# ═════════════════════════════════════════════════════════════
with tab_bericht:
    st.markdown("### Tagesbericht")
    st.caption("Erzeuge einen formalen Tagesbericht für dieses Projekt.")

    # Wetter automatisch aus der hinterlegten Baustellen-Adresse (aktueller Tag)
    wetter_vorausgefuellt = ""
    w_cached = st.session_state.get(f"wetter_{pid}", {})
    if (not w_cached or not w_cached.get("ok")) and anschrift:
        # Eigenständig laden, falls das obere Widget nicht griff
        try:
            from utils.wetter import geocode as _gc, get_wetter as _gw
            _co = st.session_state.get(f"coords_{pid}")
            if _co is None:
                _co = _gc(anschrift) or False
                st.session_state[f"coords_{pid}"] = _co
            if _co:
                w_cached = _gw(_co[0], _co[1])
                st.session_state[f"wetter_{pid}"] = w_cached
        except Exception:
            pass
    if w_cached.get("ok"):
        wetter_vorausgefuellt = (
            f"{w_cached['temp']} °C, {w_cached['beschreibung']}, "
            f"Wind {w_cached['wind']} km/h, Niederschlag {w_cached['niederschlag']} mm"
        )

    if w_cached.get("ok"):
        st.markdown(
            f"<div style='background:#f0fdf4;border:1px solid #bbf7d0;border-radius:8px;"
            f"padding:7px 13px;font-size:0.82rem;color:#166534;margin-bottom:0.7rem;'>"
            f"☀ Wetter automatisch ermittelt für <b>{html_mod.escape(anschrift)}</b> "
            f"(heute) — wird unten vorausgefüllt und ist editierbar.</div>",
            unsafe_allow_html=True
        )
    elif anschrift:
        st.markdown(
            "<div style='background:#FFF7ED;border:1px solid #fed7aa;border-radius:8px;"
            "padding:7px 13px;font-size:0.82rem;color:#9a3412;margin-bottom:0.7rem;'>"
            "Wetter konnte für die Adresse nicht ermittelt werden — bitte Wetter unten manuell eintragen "
            "oder die Baustellen-Adresse (PLZ + Ort) vervollständigen.</div>",
            unsafe_allow_html=True
        )

    with st.container(border=True):
        col_ba, col_bb = st.columns(2)
        with col_ba:
            bericht_datum = st.date_input("Datum", value=datetime.date.today(), format="DD.MM.YYYY", key=f"bericht_datum_{pid}")
        with col_bb:
            bericht_wetter = st.text_input(
                "Wetter (automatisch · aus Baustellen-Adresse)",
                value=wetter_vorausgefuellt,
                placeholder="z.B. 18 °C, bewölkt, Wind 10 km/h",
                key=f"bericht_wetter_{pid}"
            )
        col_bc, col_bd = st.columns(2)
        with col_bc:
            bericht_bauleiter = st.text_input(
                "Bauleiter",
                value=get_name(),
                placeholder="Vor- und Nachname",
                key=f"bericht_bauleiter_{pid}"
            )
        with col_bd:
            bericht_ag = st.text_input(
                "Auftraggeber",
                value=projekt.get("auftraggeber",""),
                placeholder="z.B. Stadtwerke Konstanz GmbH",
                key=f"bericht_ag_{pid}"
            )

        bericht_erledigt = st.text_area(
            "Was wurde heute erledigt?",
            height=100,
            placeholder="z.B. Erdaushub Achse A fertiggestellt, Fundamente A1–A4 betoniert...",
            key=f"bericht_erledigt_{pid}"
        )
        bericht_probleme = st.text_area(
            "Aufgetretene Probleme / Behinderungen",
            height=80,
            placeholder="z.B. Materiallieferung verzögert, Baggerausfall 2h...",
            key=f"bericht_probleme_{pid}"
        )
        bericht_personal = st.text_area(
            "Anwesende Firmen / Personal",
            height=60,
            placeholder="z.B. Erdbau GmbH (4 Mann), Zimmerei Müller (2 Mann)...",
            key=f"bericht_personal_{pid}"
        )
        bericht_fotos = st.file_uploader(
            "Foto-Anhang (optional)", type=["jpg","jpeg","png","webp"],
            accept_multiple_files=True, key=f"bericht_fotos_{pid}",
            help="Fotos von der Baustelle — werden als Anhang ins Word-Dokument eingefügt"
        )

        if st.button("Tagesbericht erstellen", type="primary",
                     disabled=not bericht_erledigt.strip(), use_container_width=True):
            with st.spinner("KI erstellt Tagesbericht..."):
                try:
                    prompt = f"""Du bist ein erfahrener Bauleiter und erstellst einen formalen Bautagesbericht.

Projekt: {pname}
Auftraggeber: {bericht_ag or 'nicht angegeben'}
Datum: {bericht_datum.strftime('%d.%m.%Y')}
Bauleiter: {bericht_bauleiter or 'nicht angegeben'}
Wetter: {bericht_wetter or 'nicht angegeben'}
Heutige Arbeiten: {bericht_erledigt}
Probleme / Behinderungen: {bericht_probleme or 'keine'}
Anwesende Firmen / Personal: {bericht_personal or 'nicht angegeben'}

Erstelle einen formalen Bautagesbericht mit:
1. Kopfdaten (Projekt, Auftraggeber, Datum, Wetter, Bauleiter)
2. Ausgeführte Arbeiten (strukturiert nach Gewerken)
3. Probleme und Behinderungen
4. Personalstärke / anwesende Firmen
5. Maßnahmen und nächste Schritte

Stil: Sachlich, formell, vollständig. Auf Deutsch."""
                    antwort = _ki_call(prompt)
                    nr = get_naechste_nummer("tagesbericht")
                    titel_tb = f"Tagesbericht Nr. {nr}"
                    speichere_dokument("tagesbericht", nr, titel_tb, antwort,
                                       {"datum": str(bericht_datum)})
                    st.session_state[f"tagesbericht_{pid}"] = {
                        "text": antwort,
                        "datum": str(bericht_datum),
                        "projekt": pname,
                        "titel": titel_tb,
                        "bauleiter": bericht_bauleiter,
                        "auftraggeber": bericht_ag,
                        "wetter": bericht_wetter,
                        "bilder": [f.getvalue() for f in bericht_fotos] if bericht_fotos else [],
                    }
                    st.rerun()
                except Exception as e:
                    err = str(e)
                    if "GEMINI_API_KEY" in err or "401" in err or "invalid_api_key" in err.lower() or "API_KEY_INVALID" in err:
                        st.error("**API-Key ungültig** — Bitte GEMINI_API_KEY in den Streamlit Secrets prüfen.")
                    elif "429" in err or "rate_limit" in err.lower():
                        st.warning("KI-Limit erreicht — Bitte 1 Minute warten und erneut versuchen.")
                    else:
                        st.error(f"KI-Fehler: {err[:300]}")

    # Neu generierten Tagesbericht anzeigen
    bericht_result = st.session_state.get(f"tagesbericht_{pid}")
    if bericht_result:
        titel_tb = bericht_result.get("titel", "Tagesbericht")
        st.markdown("---")
        st.markdown(f"#### Neu generiert: {titel_tb}")
        st.markdown(bericht_result["text"])
        st.markdown("---")
        fname_tb = titel_tb.replace(" ", "_")
        col_b1, col_b2, col_b3 = st.columns([3, 3, 1])
        with col_b1:
            st.download_button("Als .txt", data=bericht_result["text"],
                file_name=f"{fname_tb}_{bericht_result['datum']}.txt",
                mime="text/plain", use_container_width=True, key=f"tb_txt_{pid}")
        with col_b2:
            try:
                docx_b = _erstelle_docx(titel_tb,
                    {"Projekt": pname, "Datum": bericht_result["datum"],
                     "Bauleiter": bericht_result.get("bauleiter", ""),
                     "Auftraggeber": bericht_result.get("auftraggeber", ""),
                     "Wetter": bericht_result.get("wetter", "")},
                    bericht_result["text"],
                    bilder=bericht_result.get("bilder") or [])
                st.download_button("Als Word (.docx)", data=docx_b,
                    file_name=f"{fname_tb}_{bericht_result['datum']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, key=f"tb_docx_{pid}")
            except Exception as ex:
                st.error(f"Word-Fehler: {ex}")
        with col_b3:
            if st.button("Schliessen", key="del_bericht", use_container_width=True):
                del st.session_state[f"tagesbericht_{pid}"]
                st.rerun()

    # Archiv Tagesberichte
    _zeige_archiv("tagesbericht", "Tagesbericht")
